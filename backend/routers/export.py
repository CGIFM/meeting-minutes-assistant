import re
from datetime import datetime
from pathlib import Path
from urllib.parse import quote
from fastapi import APIRouter, Body
from fastapi.responses import FileResponse, Response
from pydantic import BaseModel
from typing import Optional
from db.database import get_setting

router = APIRouter()

# Obsidian vault 路径（末尾可能带空格）
OBSIDIAN_VAULT_CANDIDATES = [
    Path.home() / "Documents" / "CGIF_NOTE ",
    Path.home() / "Documents" / "CGIF_NOTE",
    Path.home() / "Documents" / "Obsidian",
]


def find_obsidian_vault() -> Optional[Path]:
    for p in OBSIDIAN_VAULT_CANDIDATES:
        if p.exists() and p.is_dir():
            return p
    # 也允许用户配置
    return None


class ExportRequest(BaseModel):
    filename: str
    content: str = ""
    minutes: str = ""
    transcript: str = ""


@router.post("/export/obsidian")
async def export_obsidian(data: ExportRequest):
    """导出到 Obsidian vault 的会议纪要文件夹。
    优先用用户设置的 obsidian_dir；否则自动检测常见 vault 路径。
    """
    user_dir = await get_setting("obsidian_dir", "")
    if user_dir and Path(user_dir).expanduser().exists():
        notes_dir = Path(user_dir).expanduser() / "会议纪要"
        notes_dir.mkdir(parents=True, exist_ok=True)
    else:
        vault = find_obsidian_vault()
        if not vault:
            return {"success": False, "message": "未配置 Obsidian 目录，且未检测到默认 vault。请在设置中填写 obsidian_dir。"}
        notes_dir = vault / "会议纪要"
        notes_dir.mkdir(parents=True, exist_ok=True)

    # 文件名清理
    safe_name = re.sub(r'[\\/:*?"<>|]', " ", data.filename)
    safe_name = safe_name.replace(".mp3", "").replace(".wav", "").replace(".m4a", "").strip()
    date_str = datetime.now().strftime("%Y-%m-%d")
    md_filename = f"{date_str}_{safe_name}.md"

    # 构造带 frontmatter 的 Obsidian 笔记
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    body = f"""---
tags: [会议纪要]
date: {now}
audio: {data.filename}
---

# {safe_name}

> 生成时间：{now}

{data.minutes}

---

## 完整转录

{data.transcript}
"""
    output_path = notes_dir / md_filename
    output_path.write_text(body, encoding="utf-8")
    return {"success": True, "message": f"已保存到 Obsidian: 会议纪要/{md_filename}", "path": str(output_path)}


@router.post("/export/pdf")
async def export_pdf(data: ExportRequest):
    """导出 PDF（使用 reportlab）"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.units import cm
        import io

        # 尝试注册中文字体（需 TrueType outline，TTC postscript outline 不支持）
        chinese_font = "Helvetica"
        for font_path in [
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/STHeiti Medium.ttc",
            "/System/Library/Fonts/Supplemental/Songti.ttc",
        ]:
            if Path(font_path).exists():
                try:
                    pdfmetrics.registerFont(TTFont("ChineseFont", font_path, subfontIndex=0))
                    chinese_font = "ChineseFont"
                    break
                except Exception:
                    continue

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
        styles = getSampleStyleSheet()
        cn_style = ParagraphStyle("CN", parent=styles["Normal"], fontName=chinese_font, fontSize=10.5, leading=18, wordWrap="CJK")
        h_style = ParagraphStyle("H", parent=styles["Heading1"], fontName=chinese_font, fontSize=16, leading=24)
        h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontName=chinese_font, fontSize=13, leading=20)

        story = []

        # 把 markdown 转成 reportlab 元素
        for line in data.minutes.split("\n"):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue
            # 简单处理 markdown
            if line.startswith("# "):
                story.append(Paragraph(line[2:].replace("*", ""), h_style))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:].replace("*", ""), h2_style))
            elif line.startswith("### "):
                story.append(Paragraph(line[4:].replace("*", ""), h2_style))
            elif line.startswith("- ") or line.startswith("* "):
                story.append(Paragraph("• " + line[2:], cn_style))
            elif line.startswith("|"):
                # 表格行简化为段落
                story.append(Paragraph(line.replace("|", " "), cn_style))
            else:
                # 加粗处理
                clean = line.replace("**", "").replace("`", "")
                story.append(Paragraph(clean, cn_style))

        doc.build(story)
        buffer.seek(0)
        safe_name = re.sub(r'[\\/:*?"<>|]', " ", data.filename).replace(".mp3", "").replace(".wav", "").replace(".m4a", "").strip()
        return Response(
            content=buffer.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(f'{safe_name}_会议纪要.pdf')}"},
        )
    except ImportError:
        return {"success": False, "message": "PDF 导出需要 reportlab，请运行 pip install reportlab"}
    except Exception as e:
        return {"success": False, "message": f"PDF 导出失败: {e}"}


@router.post("/export/word")
async def export_word(data: ExportRequest):
    """导出 Word 文档（使用 python-docx）"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        import io

        doc = Document()

        # 设置默认字体为支持中文的
        style = doc.styles["Normal"]
        style.font.name = "PingFang SC"
        style.font.size = Pt(11)

        safe_name = re.sub(r'[\\/:*?"<>|]', " ", data.filename).replace(".mp3", "").replace(".wav", "").replace(".m4a", "").strip()

        # 标题
        doc.add_heading(safe_name, 0)
        doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")

        # 解析 markdown
        for line in data.minutes.split("\n"):
            line = line.rstrip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:].replace("*", ""), level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:].replace("*", ""), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].replace("*", ""), level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:].replace("**", ""), style="List Bullet")
            elif line.startswith("|") and "---" not in line:
                # 表格简化为段落
                cells = [c.strip() for c in line.strip("|").split("|")]
                doc.add_paragraph(" | ".join(cells))
            else:
                p = doc.add_paragraph()
                # 简单加粗处理
                parts = line.split("**")
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(f'{safe_name}_会议纪要.docx')}"},
        )
    except ImportError:
        return {"success": False, "message": "Word 导出需要 python-docx，请运行 pip install python-docx"}
    except Exception as e:
        return {"success": False, "message": f"Word 导出失败: {e}"}
